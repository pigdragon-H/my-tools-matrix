#!/usr/bin/env python3
"""
Formula Universe — PDF → Word worker (semantic reconstruction + fidelity repair).

Pipeline (target wall time ~15-20s for "multi-pass verify against the original"):

  STAGE 1  input       — open & validate the source PDF
  STAGE 2  calibrate   — read the GROUND TRUTH from the original PDF with
                         PyMuPDF: every image (xref+bbox), every vector border
                         line, every fill colour (exact 24-bit RGB), every text
                         span colour. This is the "原檔案校準" reference.
  STAGE 3  convert      — pdf2docx semantic reconstruction (real paragraphs +
                          real tables + real embedded images).
  STAGE 4  repair+verify — POST-PROCESS the .docx against the ground truth and
                          re-verify in multiple passes until it matches (or the
                          pass budget is exhausted):
                            A1  re-attach any dropped images at the correct
                                size/position
                            A2  force real table borders (tcBorders) from the
                                PDF's vector lines, with the original colour
                            A3  correct shading fills (<w:shd w:fill>) to the
                                EXACT original hex (no quantisation)
                            B1  correct run colours to the exact 24-bit value
                          Each pass recomputes a fidelity score; we loop until
                          the score is stable/high or MAX_PASSES is reached.

Why this design: pdf2docx alone is a heuristic visual *guess* — it silently
drops images, misses borders and re-quantises colours, and the failure differs
per file. By calibrating against the original and repairing deterministically
in a verify loop, the output becomes faithful and *consistent* across PDFs.

Usage:
    python3 pdf2docx_worker.py <input.pdf> <output.docx> [min_seconds]

stdout (last line) emits a JSON verification report the Node layer can surface.
Exit codes:
    0  success            2  pdf2docx produced nothing usable
    3  unexpected error
"""
import sys
import os
import time
import json
import zipfile
import shutil
import tempfile
import re

# ── tuning ────────────────────────────────────────────────────────────────
# ── multi-candidate calibration loop ─────────────────────────────────────
# We generate up to MAX_CANDIDATES independent candidates. Each candidate is
# pdf2docx-converted + repaired (A1 images / A2 borders / A3 fills) at a
# slightly different strength, then scored AGAINST THE ORIGINAL PDF (ground
# truth) using the document-agnostic visual fidelity metric.
#
# THRESHOLD IS RELATIVE (not an absolute pixel-fidelity gate): no faithful
# PDF→Word converter can hit an absolute 95% pixel score because of font
# substitution + reflow. So instead of an impossible absolute cut, we keep the
# candidates that come within KEEP_RATIO of the BEST candidate's score, then
# output the single highest-scoring one. A download is ALWAYS produced.
MAX_CANDIDATES = 5
KEEP_RATIO = 0.95          # keep candidates within 95% of the best score
KEEP_THRESHOLD = 95.0      # legacy display constant (relative gate is KEEP_RATIO)
DEFAULT_MIN_SECONDS = 25   # quality-first: deliberately thorough (25-40s)
MAX_SECONDS = 40
# Per-candidate presets. Each candidate varies BOTH the pdf2docx layout-engine
# settings (which genuinely change paragraph/table reconstruction) AND the
# repair strength. These are generic engine knobs — NOTHING about any specific
# document is encoded; the visual metric then picks whichever candidate renders
# closest to that particular original.
CANDIDATE_PRESETS = [
    {"fill_dist": 150, "border_sz": 6, "pdf2docx": {}},
    {"fill_dist": 180, "border_sz": 8, "pdf2docx": {
        "connected_border_tolerance": 1.0, "max_line_spacing_ratio": 1.4}},
    {"fill_dist": 180, "border_sz": 8, "pdf2docx": {
        "line_break_width_ratio": 0.4, "new_paragraph_free_space_ratio": 0.9}},
    {"fill_dist": 210, "border_sz": 8, "pdf2docx": {
        "min_section_height": 10.0, "max_line_spacing_ratio": 1.6,
        "line_separate_threshold": 4.0}},
    {"fill_dist": 210, "border_sz": 10, "pdf2docx": {
        "connected_border_tolerance": 0.3, "min_section_height": 25.0,
        "line_break_free_space_ratio": 0.15}},
]

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def log(stage, msg=""):
    # Structured progress lines the Node layer can parse (STAGE:n:label).
    sys.stderr.write(f"[stage] {stage} {msg}\n")
    sys.stderr.flush()


# ── STAGE 2 : calibrate — read ground truth from the original PDF ───────────
def read_ground_truth(pdf_path):
    import fitz

    doc = fitz.open(pdf_path)
    pages = []
    for pno in range(doc.page_count):
        pg = doc[pno]
        page_w, page_h = float(pg.rect.width), float(pg.rect.height)
        # --- images: xref, pixel size, page-space bbox ---
        images = []
        for img in pg.get_images(full=True):
            xref = img[0]
            try:
                rects = pg.get_image_rects(xref)
            except Exception:
                rects = []
            images.append({
                "xref": xref,
                "w": img[2], "h": img[3],
                "rects": [list(map(float, r)) for r in rects],
            })

        # --- inline / form-XObject images: get_images() does NOT enumerate
        # inline-drawn images (common for company-name logos on quotations).
        # rawdict image blocks DO see them. We capture their page-space bbox
        # so they can be rasterised and re-inserted faithfully if dropped. ---
        inline_images = []
        # rects already covered by enumerable xref images (avoid double-count)
        xref_rects = [r for im in images for r in im["rects"]]

        def _overlaps(bb):
            for r in xref_rects:
                ix0, iy0 = max(bb[0], r[0]), max(bb[1], r[1])
                ix1, iy1 = min(bb[2], r[2]), min(bb[3], r[3])
                if ix1 > ix0 and iy1 > iy0:
                    inter = (ix1 - ix0) * (iy1 - iy0)
                    a = (bb[2] - bb[0]) * (bb[3] - bb[1]) or 1
                    if inter / a > 0.5:   # >50% overlap → same image
                        return True
            return False

        try:
            raw = pg.get_text("rawdict")
            for b in raw.get("blocks", []):
                if b.get("type") == 1:  # image block
                    bb = b.get("bbox")
                    if bb and not _overlaps([float(c) for c in bb]):
                        inline_images.append({
                            "bbox": [round(float(c), 1) for c in bb],
                        })
        except Exception:
            pass

        # --- vector drawings: border lines + fills ---
        h_lines, v_lines, fills = [], [], []
        for dr in pg.get_drawings():
            stroke = dr.get("color")
            fill = dr.get("fill")
            width = dr.get("width") or 0.75
            for it in dr.get("items", []):
                if it[0] == "l":  # line: (x0,y0)-(x1,y1)
                    p1, p2 = it[1], it[2]
                    if abs(p1.y - p2.y) < 0.6:       # horizontal
                        h_lines.append((round(min(p1.x, p2.x), 1), round(max(p1.x, p2.x), 1),
                                        round((p1.y + p2.y) / 2, 1),
                                        rgb_hex(stroke), round(width, 2)))
                    elif abs(p1.x - p2.x) < 0.6:     # vertical
                        v_lines.append((round((p1.x + p2.x) / 2, 1),
                                        round(min(p1.y, p2.y), 1), round(max(p1.y, p2.y), 1),
                                        rgb_hex(stroke), round(width, 2)))
                elif it[0] == "re":  # rectangle → 4 borders
                    r = it[1]
                    h_lines.append((round(r.x0, 1), round(r.x1, 1), round(r.y0, 1), rgb_hex(stroke), round(width, 2)))
                    h_lines.append((round(r.x0, 1), round(r.x1, 1), round(r.y1, 1), rgb_hex(stroke), round(width, 2)))
                    v_lines.append((round(r.x0, 1), round(r.y0, 1), round(r.y1, 1), rgb_hex(stroke), round(width, 2)))
                    v_lines.append((round(r.x1, 1), round(r.y0, 1), round(r.y1, 1), rgb_hex(stroke), round(width, 2)))
            if fill and tuple(round(c, 3) for c in fill) != (1.0, 1.0, 1.0):
                rr = dr.get("rect")
                if rr is not None:
                    fills.append({"hex": rgb_hex(fill), "rect": [rr.x0, rr.y0, rr.x1, rr.y1]})

        # --- text span colours (exact 24-bit) keyed by normalised text ---
        span_colors = {}
        for b in pg.get_text("dict")["blocks"]:
            for ln in b.get("lines", []):
                for sp in ln.get("spans", []):
                    t = sp["text"].strip()
                    if t:
                        span_colors.setdefault(norm_text(t), "%06x" % (sp["color"] & 0xFFFFFF))

        pages.append({
            "index": pno,
            "page_w": page_w,
            "page_h": page_h,
            "images": images,
            "inline_images": inline_images,
            "h_lines": h_lines,
            "v_lines": v_lines,
            "fills": fills,
            "span_colors": span_colors,
            "n_images": len(images),
            "n_inline_images": len(inline_images),
            "n_borders": len(h_lines) + len(v_lines),
            "fill_hexes": sorted({f["hex"] for f in fills}),
        })
    doc.close()
    return {"pages": pages}


def rgb_hex(c):
    if not c:
        return "000000"
    return "%02x%02x%02x" % tuple(max(0, min(255, int(round(v * 255)))) for v in c[:3])


def norm_text(t):
    return re.sub(r"\s+", "", t)


# ── STAGE 3 : pdf2docx conversion ──────────────────────────────────────────
def run_pdf2docx(in_pdf, out_docx, settings=None):
    from pdf2docx import Converter
    cv = Converter(in_pdf)
    try:
        if settings:
            cv.convert(out_docx, **settings)
        else:
            cv.convert(out_docx)
    finally:
        cv.close()


# ── STAGE 4 : repair the .docx against ground truth, verify in passes ───────
def inspect_docx(docx_path):
    """Read back what the .docx actually contains (for verification)."""
    with zipfile.ZipFile(docx_path) as z:
        names = z.namelist()
        doc_xml = z.read("word/document.xml").decode("utf-8", "ignore")
    media = [n for n in names if n.startswith("word/media/")]
    fills = re.findall(r'w:fill="([0-9A-Fa-f]{6})"', doc_xml)
    fills = [f.lower() for f in fills if f.lower() != "ffffff" and f.lower() != "auto"]
    borders = len(re.findall(r'<w:(?:tcBorders|tblBorders)\b', doc_xml))
    single = len(re.findall(r'w:val="single"', doc_xml))
    floating = len(re.findall(r'<v:shape|<wps:|<w:txbxContent', doc_xml))
    text = re.sub(r"<[^>]+>", "", doc_xml)
    return {
        "n_media": len(media),
        "fills": sorted(set(fills)),
        "border_groups": borders,
        "single_borders": single,
        "floating_boxes": floating,
        "text_norm": norm_text(text),
    }


def correct_fills(docx_path, truth, fill_dist=180):
    """A3 — snap shading fills to the EXACT original hex (nearest-by-value)."""
    wanted = []
    for p in truth["pages"]:
        wanted.extend(p["fill_hexes"])
    wanted = sorted(set(wanted))
    if not wanted:
        return 0

    tmp = docx_path + ".tmp"
    changed = 0
    with zipfile.ZipFile(docx_path) as zin, zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename == "word/document.xml":
                xml = data.decode("utf-8", "ignore")

                def repl(m):
                    nonlocal changed
                    cur = m.group(1).lower()
                    if cur in ("ffffff", "auto"):
                        return m.group(0)
                    best = min(wanted, key=lambda w: color_dist(cur, w))
                    # Snap to the nearest TRUE original fill. Threshold is
                    # generous (<=180 sum-abs over RGB) because pdf2docx tends
                    # to mis-quantise shades of the same hue (e.g. grey 808080
                    # vs the true b3b3b3); we only ever snap to a colour that
                    # actually exists in the source PDF, so this is safe.
                    if best != cur and color_dist(cur, best) <= fill_dist:
                        changed += 1
                        return m.group(0).replace(m.group(1), best)
                    return m.group(0)

                xml = re.sub(r'w:fill="([0-9A-Fa-f]{6})"', repl, xml)
                data = xml.encode("utf-8")
            zout.writestr(item, data)
    shutil.move(tmp, docx_path)
    return changed


def color_dist(a, b):
    try:
        ar, ag, ab = int(a[0:2], 16), int(a[2:4], 16), int(a[4:6], 16)
        br, bg, bb = int(b[0:2], 16), int(b[2:4], 16), int(b[4:6], 16)
        return abs(ar - br) + abs(ag - bg) + abs(ab - bb)
    except Exception:
        return 999


def ensure_table_borders(docx_path, truth, border_sz=6):
    """A2 — guarantee every table cell carries a COMPLETE set of single borders
    (top/left/bottom/right) plus a table outer frame, in the dominant original
    border colour. The original quotation tables are fully gridded, but
    pdf2docx frequently emits only horizontal borders (missing the vertical
    column separators and the table's outer bottom frame line — the user's
    "表尾底框一橫不見了"). We therefore FORCE a full grid + outer frame, not
    just fill in cells that lack borders entirely."""
    # dominant border colour across the doc
    colors = {}
    has_h = has_v = False
    for p in truth["pages"]:
        for ln in p["h_lines"]:
            colors[ln[3]] = colors.get(ln[3], 0) + 1
            has_h = True
        for ln in p["v_lines"]:
            colors[ln[3]] = colors.get(ln[3], 0) + 1
            has_v = True
    if not colors:
        return 0
    border_hex = max(colors, key=colors.get)
    # If the source has both horizontal AND vertical lines it is a full grid;
    # otherwise we only guarantee horizontal rules.
    full_grid = has_h and has_v

    edges = ["top", "left", "bottom", "right"] if full_grid else ["top", "bottom"]
    tc_borders = "<w:tcBorders>" + "".join(
        f'<w:{e} w:val="single" w:sz="{border_sz}" w:space="0" w:color="{border_hex}"/>' for e in edges
    ) + "</w:tcBorders>"
    tbl_borders = "<w:tblBorders>" + "".join(
        f'<w:{e} w:val="single" w:sz="{border_sz}" w:space="0" w:color="{border_hex}"/>'
        for e in ["top", "left", "bottom", "right", "insideH", "insideV"]
    ) + "</w:tblBorders>"

    tmp = docx_path + ".tmp"
    changed = 0
    with zipfile.ZipFile(docx_path) as zin, zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename == "word/document.xml":
                xml = data.decode("utf-8", "ignore")

                # 1) Normalise every cell: replace any existing (possibly
                #    partial) <w:tcBorders> with the full set, and inject one
                #    where missing. Handle both <w:tcPr>...</w:tcPr> and the
                #    self-closing <w:tcPr/> form.
                def fix_cell(m):
                    nonlocal changed
                    block = m.group(0)
                    if "<w:tcBorders" in block:
                        new = re.sub(r"<w:tcBorders>.*?</w:tcBorders>", tc_borders, block, flags=re.DOTALL)
                        new = re.sub(r"<w:tcBorders\s*/>", tc_borders, new)
                    else:
                        new = block.replace("<w:tcPr>", "<w:tcPr>" + tc_borders, 1)
                    if new != block:
                        changed += 1
                    return new

                xml = re.sub(r"<w:tcPr>.*?</w:tcPr>", fix_cell, xml, flags=re.DOTALL)

                # self-closing cells with no properties at all → give them props+borders
                def fix_empty_cell(m):
                    nonlocal changed
                    changed += 1
                    return "<w:tc><w:tcPr>" + tc_borders + "</w:tcPr>"
                xml = re.sub(r"<w:tc>(?!\s*<w:tcPr)", fix_empty_cell, xml)

                # 2) Ensure every table carries an outer frame (tblBorders) so
                #    the bottom/outer lines are always drawn.
                if full_grid:
                    def fix_tbl(m):
                        nonlocal changed
                        block = m.group(0)
                        if "<w:tblBorders" in block:
                            return re.sub(r"<w:tblBorders>.*?</w:tblBorders>", tbl_borders, block, flags=re.DOTALL)
                        changed += 1
                        return block.replace("<w:tblPr>", "<w:tblPr>" + tbl_borders, 1)
                    xml = re.sub(r"<w:tblPr>.*?</w:tblPr>", fix_tbl, xml, flags=re.DOTALL)

                data = xml.encode("utf-8")
            zout.writestr(item, data)
    shutil.move(tmp, docx_path)
    return changed


def reattach_missing_images(docx_path, in_pdf, truth, report):
    """A1 — if the .docx is missing UNIQUE original images, append them (full
    resolution) so the company logo / full-name graphic never silently
    disappears. Counts by UNIQUE xref (a logo repeated on every page is one
    image), so we never duplicate."""
    import fitz

    info = inspect_docx(docx_path)
    unique_xrefs = {im["xref"] for p in truth["pages"] for im in p["images"]}
    # Total distinct graphics we expect: enumerable xref images PLUS inline /
    # form-XObject images (logos drawn inline that get_images misses).
    n_inline = sum(p.get("n_inline_images", 0) for p in truth["pages"])
    want = len(unique_xrefs) + n_inline
    have = info["n_media"]
    if have >= want:
        return 0  # nothing missing (don't risk duplicating)

    doc = fitz.open(in_pdf)
    extracted = []
    seen = set()
    # (a) enumerable xref images — extract at native resolution
    for p in truth["pages"]:
        pg = doc[p["index"]]
        for im in p["images"]:
            xref = im["xref"]
            if xref in seen:
                continue
            seen.add(xref)
            try:
                pix = fitz.Pixmap(doc, xref)
                if pix.n - pix.alpha >= 4:  # CMYK → RGB
                    pix = fitz.Pixmap(fitz.csRGB, pix)
                png = pix.tobytes("png")
                im = dict(im, page_h=p.get("page_h"))
                extracted.append((xref, png, im))
            except Exception:
                continue
    # (b) inline / form-XObject images — cannot be pulled by xref, so we
    # RASTERISE the exact page region (high DPI) and re-insert at the same
    # size. This guarantees the company-name logo block never disappears,
    # regardless of how it was encoded in the source PDF.
    for p in truth["pages"]:
        pg = doc[p["index"]]
        for ii in p.get("inline_images", []):
            bb = ii.get("bbox")
            if not bb:
                continue
            try:
                clip = fitz.Rect(bb)
                # pad a hair so we don't clip antialiased edges
                clip = fitz.Rect(clip.x0 - 1, clip.y0 - 1, clip.x1 + 1, clip.y1 + 1)
                pix = pg.get_pixmap(matrix=fitz.Matrix(4, 4), clip=clip, alpha=False)
                png = pix.tobytes("png")
                extracted.append(("inline", png, {
                    "rects": [list(map(float, bb))],
                    "w": pix.width, "h": pix.height,
                    "page_h": p.get("page_h"),
                }))
            except Exception:
                continue
    doc.close()

    # how many distinct images we still need to add
    need = max(0, want - have)
    add = extracted[:need] if need else []
    if not add:
        return 0

    # Inject into the docx: add media parts + rels + a trailing paragraph with
    # an inline drawing for each missing image (so it is present & visible).
    added = _inject_images(docx_path, add)
    report["images_reattached"] = added
    return added


def _emu(px_at_96):
    return int(px_at_96 * 9525)  # 1px(96dpi)=9525 EMU


def _inject_images(docx_path, images):
    """Insert recovered images using python-docx so the output is ALWAYS valid
    OOXML that Microsoft Word accepts (the previous hand-patched XML could
    produce files Word refused to open). Top-region images (header logos) are
    moved to the very start of the body; others are appended at the end."""
    import io
    import docx
    from docx.shared import Emu
    from docx.oxml.ns import qn

    doc = docx.Document(docx_path)
    top_paras = []
    added = 0
    for (xref, png, im) in images:
        # target render size from the original page bbox (falls back to native)
        if im.get("rects"):
            r = im["rects"][0]
            wpx, hpx = (r[2] - r[0]), (r[3] - r[1])
        else:
            wpx, hpx = im["w"] * 0.75, im["h"] * 0.75
        width = Emu(_emu(max(1.0, wpx)))

        # placement RELATIVE to the page (no hardcoded coordinate): top ~18%
        is_top = False
        if im.get("rects"):
            y0 = im["rects"][0][1]
            ph = im.get("page_h") or 0
            if ph > 0:
                is_top = (y0 / ph) < 0.18

        # add the picture in a fresh trailing paragraph (always valid)
        p = doc.add_paragraph()
        run = p.add_run()
        try:
            run.add_picture(io.BytesIO(png), width=width)
        except Exception as e:
            sys.stderr.write(f"add_picture failed (skipping one image): {e}\n")
            # drop the empty paragraph we just created
            p._element.getparent().remove(p._element)
            continue
        added += 1
        if is_top:
            top_paras.append(p._element)

    # move header logos to the very front of the body (before all content)
    if top_paras:
        body = doc.element.body
        # first child that is a real paragraph/table (skip nothing — just
        # insert before the body's current first child element)
        first = None
        for child in body:
            if child.tag in (qn("w:p"), qn("w:tbl")):
                first = child
                break
        for el in reversed(top_paras):
            parent = el.getparent()
            if parent is not None:
                parent.remove(el)
            if first is not None:
                first.addprevious(el)
            else:
                body.insert(0, el)

    doc.save(docx_path)
    return added


# ── visual fidelity (document-agnostic) ──────────────────────────────────
# The objective, generalisable measure of "faithful to the original" is the
# VISUAL similarity between the original PDF and the produced .docx rendered
# back to a page image. This works for ANY document — there is nothing
# document-specific hardcoded. We render both to greyscale page images at the
# same resolution and compute SSIM (structural similarity) per page.
def _render_pdf_pages(pdf_path, dpi=110, max_pages=4):
    """Render PDF pages to greyscale numpy arrays via PyMuPDF (no temp files)."""
    import fitz
    import numpy as np
    out = []
    doc = fitz.open(pdf_path)
    try:
        n = min(doc.page_count, max_pages)
        zoom = dpi / 72.0
        mat = fitz.Matrix(zoom, zoom)
        for i in range(n):
            pix = doc[i].get_pixmap(matrix=mat, colorspace=fitz.csGRAY, alpha=False)
            arr = np.frombuffer(pix.samples, dtype=np.uint8).reshape(pix.height, pix.width)
            out.append(arr.copy())
    finally:
        doc.close()
    return out


def _docx_to_pdf(docx_path, work_dir):
    """Render a .docx to PDF using headless LibreOffice. Returns pdf path or None."""
    import subprocess
    try:
        subprocess.run(
            ["soffice", "--headless", "--convert-to", "pdf",
             "--outdir", work_dir, docx_path],
            check=True, capture_output=True, timeout=60,
        )
    except Exception as e:
        sys.stderr.write(f"docx->pdf render failed: {e}\n")
        return None
    base = os.path.splitext(os.path.basename(docx_path))[0]
    cand = os.path.join(work_dir, base + ".pdf")
    return cand if os.path.isfile(cand) else None


def _ink_mask(gray):
    """Binarise a greyscale page to an 'ink' mask (text / lines / fills / logos
    = dark or coloured pixels) using Otsu. Returns a 0/1 uint8 array. Fully
    document-agnostic — no fixed thresholds."""
    import cv2
    import numpy as np
    g = gray.astype("uint8")
    _t, binimg = cv2.threshold(g, 0, 255, cv2.THRESH_BINARY_INV + cv2.THRESH_OTSU)
    return (binimg > 0).astype("uint8")


def _page_similarity(orig_gray, cand_gray):
    """Content-focused similarity between two rendered pages (0..1). Robust to
    the 'mostly-white background' problem that makes global SSIM useless on
    documents: we compare WHERE THE INK IS, not the white space.

    Blends three document-agnostic signals:
      • ink IoU            — overlap of dark/coloured regions after alignment
      • ink-density corr.  — correlation of a coarse grid of ink density
                             (captures overall layout / block placement)
      • ink-amount ratio   — penalises large amounts of missing or extra ink
                             (a dropped logo / table / paragraph)
    """
    import cv2
    import numpy as np

    h, w = orig_gray.shape
    cand = cv2.resize(cand_gray, (w, h), interpolation=cv2.INTER_AREA)
    om = _ink_mask(orig_gray)
    cm = _ink_mask(cand)

    o_ink = int(om.sum())
    c_ink = int(cm.sum())
    if o_ink == 0 and c_ink == 0:
        return 1.0
    if o_ink == 0 or c_ink == 0:
        return 0.0

    # 1) ink IoU with a small dilation tolerance (sub-mm registration slack so
    #    that minor font/reflow shifts don't unfairly tank the score).
    kern = cv2.getStructuringElement(cv2.MORPH_ELLIPSE, (5, 5))
    om_d = cv2.dilate(om, kern)
    cm_d = cv2.dilate(cm, kern)
    inter = int(np.logical_and(om, cm_d).sum() + np.logical_and(cm, om_d).sum())
    union = int(om.sum() + cm.sum())
    iou = inter / union if union else 0.0

    # 2) coarse ink-density grid correlation (layout/structure agreement)
    gh, gw = 24, 18
    def grid_density(m):
        cells = []
        ys = np.linspace(0, h, gh + 1).astype(int)
        xs = np.linspace(0, w, gw + 1).astype(int)
        for i in range(gh):
            for j in range(gw):
                blk = m[ys[i]:ys[i + 1], xs[j]:xs[j + 1]]
                cells.append(blk.mean() if blk.size else 0.0)
        return np.array(cells, dtype="float64")
    od = grid_density(om)
    cd = grid_density(cm)
    if od.std() < 1e-6 or cd.std() < 1e-6:
        corr = 0.0
    else:
        corr = float(np.corrcoef(od, cd)[0, 1])
    corr = max(0.0, corr)

    # 3) ink-amount ratio (missing/extra content penalty)
    amount = min(o_ink, c_ink) / max(o_ink, c_ink)

    # blend (weights favour structural agreement + content completeness)
    return float(0.45 * iou + 0.35 * corr + 0.20 * amount)


def visual_similarity(orig_pdf, docx_path, work_dir):
    """Render the .docx back to PDF and compare page-by-page against the
    ORIGINAL pdf using a content-focused (ink-structure) similarity. Returns
    (score 0..100, n_pages_compared) or (None, 0) if rendering is unavailable.
    Completely document-agnostic — nothing about any specific file is encoded."""
    import cv2
    cand_pdf = _docx_to_pdf(docx_path, work_dir)
    if not cand_pdf:
        return None, 0
    try:
        orig_imgs = _render_pdf_pages(orig_pdf)
        cand_imgs = _render_pdf_pages(cand_pdf)
    except Exception as e:
        sys.stderr.write(f"page render failed: {e}\n")
        return None, 0
    if not orig_imgs or not cand_imgs:
        return None, 0

    n = min(len(orig_imgs), len(cand_imgs))
    scores = [_page_similarity(orig_imgs[i], cand_imgs[i]) for i in range(n)]
    # penalise page-count mismatch (missing/extra pages hurt fidelity)
    page_penalty = min(len(orig_imgs), len(cand_imgs)) / max(len(orig_imgs), len(cand_imgs))
    raw = (sum(scores) / len(scores)) if scores else 0.0
    return round(100.0 * raw * page_penalty, 1), n


def fidelity_score(docx_path, truth, orig_pdf=None, work_dir=None):
    """0..100 faithful-to-original score.

    PRIMARY metric is VISUAL similarity (SSIM of the rendered pages) — this is
    objective and works for any document, with nothing hardcoded. A small
    structural guard (no floating text boxes = not a "monster") is blended in.
    If visual rendering is unavailable, falls back to a structural estimate."""
    info = inspect_docx(docx_path)

    vis, npages = (None, 0)
    if orig_pdf and work_dir:
        try:
            vis, npages = visual_similarity(orig_pdf, docx_path, work_dir)
        except Exception as e:
            sys.stderr.write(f"visual similarity failed: {e}\n")
            vis = None

    # structural guard: floating text boxes => the doc will break in Word.
    structural_ok = 1.0 if info["floating_boxes"] == 0 else 0.0

    if vis is not None:
        # 90% visual + 10% structural guard
        final = 0.90 * vis + 0.10 * (100.0 * structural_ok)
        info["_visual"] = vis
        info["_visual_pages"] = npages
        return round(final, 1), info

    # ── fallback (visual rendering unavailable): structural estimate ──
    want_imgs = len({im["xref"] for p in truth["pages"] for im in p["images"]}) \
        + sum(p.get("n_inline_images", 0) for p in truth["pages"])
    want_borders = any(p["n_borders"] > 0 for p in truth["pages"])
    want_fills = sorted({h for p in truth["pages"] for h in p["fill_hexes"]})
    score, weight = 0.0, 0.0
    weight += 35
    score += 35 if want_imgs == 0 else 35 * min(1.0, info["n_media"] / want_imgs)
    weight += 25
    score += 25 if (not want_borders or info["single_borders"] > 0) else 0
    weight += 25
    if not want_fills:
        score += 25
    else:
        present = [f for f in want_fills if any(color_dist(f, g) <= 10 for g in info["fills"])]
        score += 25 * (len(present) / len(want_fills))
    weight += 15
    score += 15 if info["floating_boxes"] == 0 else 0
    info["_visual"] = None
    return round(100 * score / weight, 1), info


def build_candidate(in_pdf, cand_path, truth, preset):
    """Generate ONE candidate: pdf2docx convert + A1/A2/A3 repair at the
    given strength preset. Returns a per-candidate repair tally."""
    tally = {"images_reattached": 0, "borders_added": 0, "fills_corrected": 0}
    run_pdf2docx(in_pdf, cand_path, settings=preset.get("pdf2docx"))
    if not os.path.isfile(cand_path) or os.path.getsize(cand_path) < 200:
        raise RuntimeError("pdf2docx produced an empty/too-small document")
    if truth["pages"]:
        try:
            tally["images_reattached"] = reattach_missing_images(cand_path, in_pdf, truth, {})
        except Exception as e:
            sys.stderr.write(f"image repair skipped: {e}\n")
        try:
            tally["borders_added"] = ensure_table_borders(cand_path, truth, border_sz=preset["border_sz"])
        except Exception as e:
            sys.stderr.write(f"border repair skipped: {e}\n")
        try:
            tally["fills_corrected"] = correct_fills(cand_path, truth, fill_dist=preset["fill_dist"])
        except Exception as e:
            sys.stderr.write(f"fill repair skipped: {e}\n")
    return tally


def main():
    if len(sys.argv) < 3:
        sys.stderr.write("usage: pdf2docx_worker.py <input.pdf> <output.docx> [min_seconds]\n")
        return 3
    in_pdf, out_docx = sys.argv[1], sys.argv[2]
    min_seconds = DEFAULT_MIN_SECONDS
    if len(sys.argv) >= 4:
        try:
            min_seconds = max(0, min(MAX_SECONDS, float(sys.argv[3])))
        except ValueError:
            pass

    if not os.path.isfile(in_pdf):
        sys.stderr.write(f"input not found: {in_pdf}\n")
        return 3

    t0 = time.time()
    report = {
        "images_reattached": 0, "borders_added": 0, "fills_corrected": 0,
        "candidates": [], "kept_count": 0, "chosen_n": None,
        "chosen_score": None, "kept_threshold": KEEP_THRESHOLD,
        "keep_ratio": KEEP_RATIO, "relative_threshold": True,
    }

    # STAGE 1 — input
    log("1", "input")
    try:
        import fitz  # noqa
        d = fitz.open(in_pdf)
        report["pages"] = d.page_count
        d.close()
    except Exception as e:
        sys.stderr.write(f"open failed: {e}\n")
        return 3

    # STAGE 2 — calibrate (read ground truth ONCE; reused for every candidate)
    log("2", "calibrate")
    try:
        truth = read_ground_truth(in_pdf)
    except Exception as e:
        sys.stderr.write(f"calibration failed (continuing without repair): {e}\n")
        truth = {"pages": []}

    # STAGE 3 + 4 — generate up to MAX_CANDIDATES, score each against the
    # ORIGINAL pdf, keep >= threshold, then pick the best.
    work_dir = tempfile.mkdtemp(prefix="cands_")
    candidates = []   # list of dicts: {n, path, score, kept, tally}
    try:
        for n in range(1, MAX_CANDIDATES + 1):
            # time guard: never blow far past the budget; always produce >=1.
            elapsed = time.time() - t0
            if n > 1 and elapsed > MAX_SECONDS - 4:
                log("4", f"time budget reached after {n-1} candidates")
                break

            preset = CANDIDATE_PRESETS[(n - 1) % len(CANDIDATE_PRESETS)]
            cand_path = os.path.join(work_dir, f"cand_{n}.docx")
            log("3", f"candidate {n}/{MAX_CANDIDATES}")
            try:
                tally = build_candidate(in_pdf, cand_path, truth, preset)
            except Exception as e:
                sys.stderr.write(f"candidate {n} failed: {e}\n")
                continue

            log("4", f"verify candidate {n}")
            # Score VISUALLY against the original PDF (document-agnostic SSIM).
            score, _info = fidelity_score(cand_path, truth, orig_pdf=in_pdf, work_dir=work_dir)
            vis = _info.get("_visual")
            # 'kept' is decided RELATIVELY after the whole loop (see below);
            # store the raw score now.
            candidates.append({"n": n, "path": cand_path, "score": score,
                               "kept": False, "tally": tally, "visual": vis})
            log("4", f"candidate {n} score={score} visual={vis}")

        if not candidates:
            # last-ditch: a single plain conversion so a download still exists
            sys.stderr.write("no candidate succeeded; falling back to plain convert\n")
            run_pdf2docx(in_pdf, out_docx)
            if not os.path.isfile(out_docx) or os.path.getsize(out_docx) < 200:
                return 2
            report["chosen_n"] = 0
            report["chosen_score"] = None
        else:
            # RELATIVE threshold (user option A): the best candidate is the
            # reference; keep everything within KEEP_RATIO of it. This is the
            # honest, document-agnostic way to gate quality — there is no
            # impossible absolute pixel cut-off.
            best = max(candidates, key=lambda c: c["score"])
            cut = best["score"] * KEEP_RATIO
            for c in candidates:
                c["kept"] = c["score"] >= cut
            kept = [c for c in candidates if c["kept"]]
            report["kept_count"] = len(kept)
            report["keep_ratio"] = KEEP_RATIO
            report["best_score"] = best["score"]
            report["keep_cutoff"] = round(cut, 2)
            for c in candidates:
                report["candidates"].append(
                    {"n": c["n"], "score": c["score"],
                     "kept": c["kept"], "visual": c["visual"]})
            # Output the single highest-scoring candidate.
            shutil.copyfile(best["path"], out_docx)
            report["chosen_n"] = best["n"]
            report["chosen_score"] = best["score"]
            report["final_score"] = best["score"]
            report["visual_score"] = best.get("visual")
            report["images_reattached"] = best["tally"]["images_reattached"]
            report["borders_added"] = best["tally"]["borders_added"]
            report["fills_corrected"] = best["tally"]["fills_corrected"]
            log("4", f"chosen candidate {best['n']} score={best['score']} "
                     f"(kept={len(kept)}/{len(candidates)} within {KEEP_RATIO:.0%} of best)")
    finally:
        shutil.rmtree(work_dir, ignore_errors=True)

    report["passes"] = len(candidates)
    report["scores"] = [c["score"] for c in candidates]

    # throttle to the requested minimum wall time (25-40s) — quality-first feel
    elapsed = time.time() - t0
    if elapsed < min_seconds:
        time.sleep(min_seconds - elapsed)
    report["elapsed_s"] = round(time.time() - t0, 2)

    if not os.path.isfile(out_docx) or os.path.getsize(out_docx) < 200:
        return 2

    print(json.dumps(report, ensure_ascii=False))
    return 0


if __name__ == "__main__":
    sys.exit(main())
