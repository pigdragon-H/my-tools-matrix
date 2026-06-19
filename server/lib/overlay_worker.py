#!/usr/bin/env python3
"""
Strategy-1 "text-over-image" PDF -> Word engine (prototype v5).

v5 goals (design-type PDFs: colored fills, white text, emoji):
  * Per-line text color taken from the dominant span's real color.
  * Background redaction fill is SAMPLED from the pixels just outside the
    text bbox (so white text over a dark band keeps the dark band; the
    overlaid white text is then visible).
  * Emoji characters are split into separate runs using an emoji font so
    they render (instead of tofu boxes).
  * CJK -> a serif/sans CJK font; latin -> Times New Roman.
Output: every page = full-page background image + absolutely-positioned,
selectable/editable text boxes on top. Fast, never hangs.
"""
import sys
import io
import re
from xml.sax.saxutils import escape
import fitz
from docx import Document
from docx.shared import Pt, Emu
from docx.enum.section import WD_SECTION
from docx.oxml import parse_xml

DPI = 200
EMU_PER_PT = 12700

# Emoji / pictographic ranges (common subset)
_EMOJI_RE = re.compile(
    "[" 
    "\U0001F300-\U0001FAFF"
    "\U00002600-\U000027BF"
    "\U0001F000-\U0001F0FF"
    "\U00002190-\U000021FF"
    "\U00002B00-\U00002BFF"
    "\U0000FE00-\U0000FE0F"
    "]"
)

NS = (
    'xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing" '
    'xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" '
    'xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape" '
    'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
)

_uid = [1000]


def _runs_xml(text, color_hex, cjk_font, bold):
    """Split text into emoji/non-emoji runs; emoji uses Segoe UI Emoji."""
    out = []
    b = '<w:b/>' if bold else ''
    parts = []
    last = 0
    for m in _EMOJI_RE.finditer(text):
        if m.start() > last:
            parts.append(("t", text[last:m.start()]))
        parts.append(("e", m.group()))
        last = m.end()
    if last < len(text):
        parts.append(("t", text[last:]))
    if not parts:
        parts = [("t", text)]
    for kind, seg in parts:
        if not seg:
            continue
        txt = escape(seg)
        if kind == "e":
            font = "Segoe UI Emoji"
            out.append(
                f'<w:r><w:rPr><w:rFonts w:ascii="{font}" w:hAnsi="{font}" w:eastAsia="{font}"/>'
                f'<w:sz w:val="{{SZ}}"/></w:rPr><w:t xml:space="preserve">{txt}</w:t></w:r>'
            )
        else:
            out.append(
                f'<w:r><w:rPr><w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman" '
                f'w:eastAsia="{cjk_font}"/>{b}<w:color w:val="{color_hex}"/>'
                f'<w:sz w:val="{{SZ}}"/></w:rPr><w:t xml:space="preserve">{txt}</w:t></w:r>'
            )
    return "".join(out)


def textbox_xml(text, left_pt, top_pt, width_pt, height_pt, font_pt, color_rgb, cjk_font, bold):
    _uid[0] += 1
    cx = int(max(width_pt, 2) * EMU_PER_PT)
    cy = int(max(height_pt, font_pt * 1.1) * EMU_PER_PT)
    ox = int(left_pt * EMU_PER_PT)
    oy = int(top_pt * EMU_PER_PT)
    sz = int(font_pt * 2)
    color = '%02X%02X%02X' % color_rgb
    runs = _runs_xml(text, color, cjk_font, bold).replace("{SZ}", str(sz))
    return f'''<w:drawing {NS}>
  <wp:anchor distT="0" distB="0" distL="0" distR="0" simplePos="0" relativeHeight="{_uid[0]}" behindDoc="0" locked="0" layoutInCell="1" allowOverlap="1">
    <wp:simplePos x="0" y="0"/>
    <wp:positionH relativeFrom="page"><wp:posOffset>{ox}</wp:posOffset></wp:positionH>
    <wp:positionV relativeFrom="page"><wp:posOffset>{oy}</wp:posOffset></wp:positionV>
    <wp:extent cx="{cx}" cy="{cy}"/>
    <wp:effectExtent l="0" t="0" r="0" b="0"/>
    <wp:wrapNone/>
    <wp:docPr id="{_uid[0]}" name="TB{_uid[0]}"/>
    <a:graphic>
      <a:graphicData uri="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">
        <wps:wsp>
          <wps:cNvSpPr txBox="1"/>
          <wps:spPr>
            <a:xfrm><a:off x="0" y="0"/><a:ext cx="{cx}" cy="{cy}"/></a:xfrm>
            <a:prstGeom prst="rect"><a:avLst/></a:prstGeom>
            <a:noFill/>
          </wps:spPr>
          <wps:txbx>
            <w:txbxContent>
              <w:p>
                <w:pPr><w:spacing w:before="0" w:after="0" w:line="240" w:lineRule="auto"/></w:pPr>
                {runs}
              </w:p>
            </w:txbxContent>
          </wps:txbx>
          <wps:bodyPr wrap="none" lIns="0" tIns="0" rIns="0" bIns="0" anchor="t"/>
        </wps:wsp>
      </a:graphicData>
    </a:graphic>
  </wp:anchor>
</w:drawing>'''


def extract_lines(page):
    lines = []
    for blk in page.get_text("dict").get("blocks", []):
        for line in blk.get("lines", []):
            spans = [s for s in line.get("spans", []) if s.get("text", "").strip()]
            if not spans:
                continue
            text = "".join(s["text"] for s in line["spans"])
            if not text.strip():
                continue
            x0 = min(s["bbox"][0] for s in spans)
            y0 = min(s["bbox"][1] for s in spans)
            x1 = max(s["bbox"][2] for s in spans)
            y1 = max(s["bbox"][3] for s in spans)
            dom = max(spans, key=lambda s: s["bbox"][2] - s["bbox"][0])
            c = dom.get("color", 0)
            rgb = ((c >> 16) & 255, (c >> 8) & 255, c & 255)
            max_size = max(s.get("size", 10) for s in spans)
            font = dom.get("font", "").lower()
            bold = "bold" in font or "black" in font or "heavy" in font
            lines.append({
                "text": text, "x0": x0, "y0": y0, "x1": x1, "y1": y1,
                "size": max_size, "color": rgb, "bold": bold,
            })
    return lines


def _is_lightish(rgb):
    return (rgb[0] * 0.299 + rgb[1] * 0.587 + rgb[2] * 0.114) > 160


def render_background(page, lines, base_pix):
    """Redact text bboxes, sampling fill color from just-left pixels so that
    text on colored bands keeps the band color. Light (e.g. white) text uses
    sampled neighbor color; dark text on white uses white."""
    page_pix = base_pix  # pre-rendered pixmap (no redaction) for sampling
    pw = page.rect.width
    ph = page.rect.height
    sx = page_pix.width / pw
    sy = page_pix.height / ph

    def sample(px, py):
        ix = min(max(int(px * sx), 0), page_pix.width - 1)
        iy = min(max(int(py * sy), 0), page_pix.height - 1)
        p = page_pix.pixel(ix, iy)
        return (p[0] / 255.0, p[1] / 255.0, p[2] / 255.0)

    for ln in lines:
        r = fitz.Rect(ln["x0"], ln["y0"], ln["x1"], ln["y1"])
        if r.is_empty or r.width <= 0 or r.height <= 0:
            continue
        # sample a point just left of the text, vertical middle
        sxp = max(ln["x0"] - 3, 0)
        syp = (ln["y0"] + ln["y1"]) / 2.0
        fill = sample(sxp, syp)
        # if light text, also try sampling just above to get band color
        if _is_lightish(ln["color"]):
            above = sample((ln["x0"] + ln["x1"]) / 2.0, max(ln["y0"] - 2, 0))
            fill = above
        page.add_redact_annot(r, fill=fill)
    try:
        page.apply_redactions(images=fitz.PDF_REDACT_IMAGE_NONE)
    except Exception:
        try:
            page.apply_redactions()
        except Exception:
            pass
    return page.get_pixmap(dpi=DPI).tobytes("png")


def pick_cjk_font(text):
    return "SimSun"


def build(pdf_path: str, out_path: str):
    src = fitz.open(pdf_path)
    word = Document()
    for i in range(len(src)):
        page = src[i]
        pw_pt, ph_pt = page.rect.width, page.rect.height
        lines = extract_lines(page)
        base_pix = page.get_pixmap(dpi=DPI)  # before redaction, for sampling
        bg_png = render_background(page, lines, base_pix)

        sec = word.sections[i] if i < len(word.sections) else word.add_section(WD_SECTION.NEW_PAGE)
        sec.page_width = Emu(int(pw_pt * EMU_PER_PT))
        sec.page_height = Emu(int(ph_pt * EMU_PER_PT))
        sec.top_margin = Pt(0); sec.bottom_margin = Pt(0)
        sec.left_margin = Pt(0); sec.right_margin = Pt(0)
        sec.header_distance = Pt(0); sec.footer_distance = Pt(0)

        p_bg = word.add_paragraph()
        p_bg.paragraph_format.space_after = Pt(0)
        run_bg = p_bg.add_run()
        run_bg.add_picture(io.BytesIO(bg_png), width=Emu(int(pw_pt * EMU_PER_PT)))

        for ln in lines:
            w = (ln["x1"] - ln["x0"]) * 1.35 + max(24, ln["size"] * 2)
            h = (ln["y1"] - ln["y0"]) + 2
            xml = textbox_xml(ln["text"], ln["x0"], ln["y0"], w, h,
                              ln["size"], ln["color"], pick_cjk_font(ln["text"]), ln["bold"])
            run = p_bg.add_run()
            run._r.append(parse_xml(xml))

    src.close()
    word.save(out_path)
    print("OK", out_path, "pages=", len(word.sections))


if __name__ == "__main__":
    build(sys.argv[1], sys.argv[2])
